from .base import Document
import docx


class DocxDocument(Document):
    def get_text(self):
        doc = docx.Document(self.file)
        return '\n'.join(p.text for p in doc.paragraphs)

    def generate_copy(self):
        pass
